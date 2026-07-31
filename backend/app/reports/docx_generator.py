"""Professional DOCX report generator.

Takes AI-generated markdown report content and produces a formatted
Word document with cover page, severity charts, and detailed findings.
"""
import os
import re
import logging
from datetime import datetime

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx not installed. DOCX report generation unavailable.")


SEVERITY_COLORS = {
    "critical": RGBColor(220, 38, 38) if HAS_DOCX else None,
    "high": RGBColor(249, 115, 22) if HAS_DOCX else None,
    "medium": RGBColor(234, 179, 8) if HAS_DOCX else None,
    "low": RGBColor(59, 130, 246) if HAS_DOCX else None,
    "info": RGBColor(107, 114, 128) if HAS_DOCX else None,
}

DARK_BG = RGBColor(26, 26, 37) if HAS_DOCX else None
ACCENT_RED = RGBColor(220, 38, 38) if HAS_DOCX else None
GRAY_TEXT = RGBColor(113, 113, 122) if HAS_DOCX else None
WHITE = RGBColor(255, 255, 255) if HAS_DOCX else None


def _hex_to_rgb(hex_color):
    """Convert hex color string to RGBColor."""
    if not HAS_DOCX:
        return None
    if not hex_color:
        return ACCENT_RED
    hex_color = hex_color.lstrip("#")
    try:
        r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
        return RGBColor(r, g, b)
    except (ValueError, IndexError):
        return ACCENT_RED


def generate_docx_report(
    engagement,
    findings,
    attack_paths,
    ai_content: str,
    output_path: str,
    template=None,
) -> str:
    """Generate a professional DOCX pentest report."""
    if not HAS_DOCX:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")

    # Template branding (defaults if no template)
    brand_company = (template.company_name if template and template.company_name
                     else "Breachwright | Advent Cybersecurity")
    brand_color = _hex_to_rgb(template.primary_color) if template and template.primary_color else ACCENT_RED
    brand_header = template.header_text if template and template.header_text else "PENETRATION TEST REPORT"
    brand_footer = template.footer_text if template and template.footer_text else "CONFIDENTIAL"
    brand_logo = template.logo_path if template and template.logo_path else None

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(30, 30, 30)

    # Configure heading styles
    for level, size, color in [
        ("Heading 1", 22, brand_color),
        ("Heading 2", 16, RGBColor(30, 30, 30)),
        ("Heading 3", 13, RGBColor(60, 60, 60)),
    ]:
        s = doc.styles[level]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.color.rgb = color
        s.font.bold = True

    # ---- COVER PAGE ----

    # Logo (if template has one)
    if brand_logo and os.path.exists(brand_logo):
        try:
            logo_p = doc.add_paragraph()
            logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logo_p.add_run().add_picture(brand_logo, width=Inches(2.5))
            doc.add_paragraph()
        except Exception:
            for _ in range(3):
                doc.add_paragraph()
    else:
        for _ in range(6):
            doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(brand_header)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = brand_color
    run.font.name = "Calibri"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(engagement.name)
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(60, 60, 60)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Client: {engagement.client_name}")
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY_TEXT

    if engagement.scope:
        scope_p = doc.add_paragraph()
        scope_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = scope_p.add_run(f"Scope: {engagement.scope}")
        run.font.size = Pt(11)
        run.font.color.rgb = GRAY_TEXT

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dates = []
    if engagement.start_date:
        dates.append(f"Start: {engagement.start_date}")
    if engagement.end_date:
        dates.append(f"End: {engagement.end_date}")
    if not dates:
        dates.append(f"Date: {datetime.now().strftime('%B %d, %Y')}")
    run = date_p.add_run(" | ".join(dates))
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY_TEXT

    for _ in range(4):
        doc.add_paragraph()

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(brand_footer)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = brand_color

    branding = doc.add_paragraph()
    branding.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = branding.add_run(f"Generated by {brand_company}")
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY_TEXT

    doc.add_page_break()

    # ---- SEVERITY SUMMARY TABLE ----
    doc.add_heading("Findings Summary", level=1)

    sev_counts = {"critical": 0, "high": 0, "medium": 0, "low": 0, "info": 0}
    for f in findings:
        sev = f.severity.value if hasattr(f.severity, "value") else f.severity
        sev_counts[sev] = sev_counts.get(sev, 0) + 1

    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Total", "Critical", "High", "Medium", "Low", "Info"]
    values = [
        str(len(findings)),
        str(sev_counts["critical"]),
        str(sev_counts["high"]),
        str(sev_counts["medium"]),
        str(sev_counts["low"]),
        str(sev_counts["info"]),
    ]

    hdr_row = table.rows[0]
    for i, (h, v) in enumerate(zip(headers, values)):
        cell = hdr_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{h}\n{v}")
        run.font.size = Pt(10)
        run.bold = True
        if i > 0:
            sev_key = h.lower()
            if sev_key in SEVERITY_COLORS and SEVERITY_COLORS[sev_key]:
                run.font.color.rgb = SEVERITY_COLORS[sev_key]

    doc.add_paragraph()

    # ---- FINDINGS TABLE ----
    if findings:
        ftable = doc.add_table(rows=1, cols=5)
        ftable.style = "Table Grid"
        ftable.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr = ftable.rows[0]
        for i, h in enumerate(["#", "Title", "Severity", "CVSS", "Hosts"]):
            hdr.cells[i].text = h
            for p in hdr.cells[i].paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(9)

        for idx, f in enumerate(findings, 1):
            row = ftable.add_row()
            sev = f.severity.value if hasattr(f.severity, "value") else f.severity
            row.cells[0].text = str(idx)
            row.cells[1].text = f.title or ""
            row.cells[2].text = sev.upper()
            row.cells[3].text = str(f.cvss_score) if f.cvss_score is not None else "N/A"
            row.cells[4].text = (f.affected_hosts or "N/A")[:50]

            # Color the severity cell
            for p in row.cells[2].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.bold = True
                    if sev in SEVERITY_COLORS and SEVERITY_COLORS[sev]:
                        r.font.color.rgb = SEVERITY_COLORS[sev]

            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        if r.font.size is None:
                            r.font.size = Pt(9)

    doc.add_page_break()

    # ---- AI-GENERATED CONTENT (markdown to docx) ----
    _markdown_to_docx(doc, _docx_body_markdown(ai_content))

    # Save
    doc.save(output_path)
    return output_path


def _docx_body_markdown(markdown: str) -> str:
    """Remove sections already represented by the DOCX cover and summary."""
    parts = re.split(r"(?m)(?=^## )", markdown)
    if len(parts) == 1:
        return markdown

    body = []
    preamble = parts[0].strip()
    if preamble and not preamble.startswith("# "):
        body.append(preamble)
    skipped = {"engagement overview", "findings summary"}
    for part in parts[1:]:
        heading = part.splitlines()[0][3:].strip().lower()
        if heading in skipped:
            continue
        body.append(part.strip())
    return "\n\n".join(body)


def _markdown_to_docx(doc, md_text):
    """Convert markdown text to docx paragraphs."""
    lines = md_text.split("\n")
    in_code_block = False

    for line in lines:
        stripped = line.strip()

        # Code blocks
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            continue

        if in_code_block:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(50, 50, 50)
            continue

        # Headings
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("---") or stripped.startswith("==="):
            continue  # Skip horizontal rules
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped)
            doc.add_paragraph(text, style="List Number")
        elif stripped == "":
            continue
        else:
            # Regular paragraph with basic inline formatting
            p = doc.add_paragraph()
            _add_formatted_text(p, stripped)


def _add_formatted_text(paragraph, text):
    """Add text with basic markdown inline formatting (bold, italic, code)."""
    parts = re.split(r"(\*\*.*?\*\*|`.*?`|\*.*?\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(180, 50, 50)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)
