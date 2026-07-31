import io
import os
import sys
import unittest
from datetime import date
from types import SimpleNamespace


ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
BACKEND = os.path.join(ROOT, "backend")
if BACKEND not in sys.path:
    sys.path.insert(0, BACKEND)

from app.engagements.models import Severity
from app.reports.content import build_report_content
from app.reports.docx_generator import generate_docx_report


class ReportContentTests(unittest.TestCase):
    def test_report_contains_engagement_findings_and_brand_attribution(self):
        engagement = SimpleNamespace(
            name="External Assessment",
            client_name="Example Client",
            scope="example.test",
            start_date=date(2026, 7, 1),
            end_date=date(2026, 7, 15),
        )
        finding = SimpleNamespace(
            title="Outdated service",
            severity=Severity.high,
            cvss_score=8.1,
            affected_hosts="10.0.0.5",
            retest_status="open",
            description="A service requires an update.",
            evidence="Version response captured.",
            remediation="Install the supported release.",
            evidence_refs=[],
            ai_inference=False,
        )
        attack_path = SimpleNamespace(
            name="External foothold",
            risk_level="high",
            description="The exposed service permits initial access.",
            narrative="A reviewed attack narrative.",
            mitre_techniques=[{
                "technique_id": "T1190",
                "technique_name": "Exploit Public-Facing Application",
            }],
        )

        report = build_report_content(engagement, [finding], [attack_path])

        self.assertIn("External Assessment - Penetration Test Report", report)
        self.assertIn("created by Advent Cybersecurity", report)
        self.assertIn("| High | 1 |", report)
        self.assertIn("Outdated service", report)
        self.assertIn("Install the supported release.", report)
        self.assertIn("External foothold", report)
        self.assertIn("A reviewed attack narrative.", report)
        self.assertIn("T1190: Exploit Public-Facing Application", report)

    def test_docx_preserves_zero_cvss_and_does_not_duplicate_attack_paths(self):
        engagement = SimpleNamespace(
            name="Zero Score Assessment",
            client_name="Example Client",
            scope=None,
            start_date=None,
            end_date=None,
        )
        finding = SimpleNamespace(
            title="Informational observation",
            severity=Severity.info,
            cvss_score=0.0,
            affected_hosts=None,
            retest_status="open",
            description="No direct security impact.",
            evidence="Observed during review.",
            remediation="No action required.",
            evidence_refs=[],
            ai_inference=False,
        )
        attack_path = SimpleNamespace(
            name="Reviewed chain",
            risk_level="low",
            description="A reviewed chain.",
            narrative=None,
            mitre_techniques=[],
        )
        content = build_report_content(engagement, [finding], [attack_path])

        output = io.BytesIO()
        generate_docx_report(
            engagement,
            [finding],
            [attack_path],
            content,
            output,
        )
        output.seek(0)
        from docx import Document

        document = Document(output)
        paragraphs = "\n".join(p.text for p in document.paragraphs)
        table_text = "\n".join(
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        )

        self.assertIn("0.0", table_text)
        self.assertNotIn("CVSS\nN/A", table_text)
        self.assertEqual(paragraphs.count("Findings Summary"), 1)
        self.assertNotIn("Engagement Overview", paragraphs)
        self.assertEqual(paragraphs.count("Executive Summary"), 1)
        self.assertEqual(paragraphs.count("Reviewed chain"), 1)

    def test_empty_engagement_still_produces_complete_report(self):
        engagement = SimpleNamespace(
            name="Empty Assessment",
            client_name="Example Client",
            scope=None,
            start_date=None,
            end_date=None,
        )

        report = build_report_content(engagement, [], [])

        self.assertIn("No findings were recorded", report)
        self.assertIn("No attack paths were recorded", report)
        self.assertIn("## Conclusion", report)


if __name__ == "__main__":
    unittest.main()
